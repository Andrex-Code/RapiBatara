from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\pipev\OneDrive\Documentos\RapiBatará\rapibatara_especificacion_tecnica_v1.docx")

ACCENT = "2E74B5"
DARK = "1F1F1F"
MUTED = "5F6B7A"
LIGHT = "F4F6F9"
DARK_BLUE = "1F4D78"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("left", "top", "right", "bottom"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def remove_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def fmt_run(run, size=11, color=DARK, bold=False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def fmt_style(style, size, color, bold=True, before=0, after=0, line=1.15) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def add_para(doc, text="", *, style=None, size=11, color=DARK, bold=False, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    fmt_run(r, size=size, color=color, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        fmt_run(r, size=11)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        fmt_run(r, size=11)


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(1.75), Inches(4.75)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    for r_idx, (left, right) in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate((left, right)):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_margins(cell)
            if c_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(value)
            fmt_run(r, size=10.5, bold=(c_idx == 0))
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = width

    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_margins(cell)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(header)
        fmt_run(r, size=10.5, bold=True)

    for r_idx, rowvals in enumerate(rows, start=1):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(rowvals):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            fmt_run(r, size=10)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    fmt_run(r1, size=11, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    fmt_run(r2, size=10.5)
    remove_borders(table)
    return table


def format_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    fmt_style(normal, size=11, color=DARK, bold=False, after=6, line=1.25)
    normal.paragraph_format.space_before = Pt(0)

    fmt_style(doc.styles["Heading 1"], size=16, color=ACCENT, bold=True, before=18, after=10, line=1.15)
    fmt_style(doc.styles["Heading 2"], size=13, color=ACCENT, bold=True, before=14, after=7, line=1.15)
    fmt_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=10, after=5, line=1.15)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        fmt_style(style, size=11, color=DARK, bold=False, after=4, line=1.15)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("RapiBatará | Especificación técnica y funcional")
    fmt_run(r, size=9, color=MUTED)


def build_document():
    doc = Document()
    format_doc(doc)
    add_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("RapiBatará")
    fmt_run(r, size=26, color="000000", bold=False)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("Especificación funcional y técnica para una app de tienda de conjunto cerrado")
    fmt_run(r2, size=11, color=MUTED)

    add_callout(
        doc,
        "Decisión de producto:",
        "La solución recomendada es una app responsive tipo PWA, pensada para vecinos, tendero y domiciliario en un solo sistema. Debe resolver compras, comprobación de pagos, inventario, despacho interno y control operativo con trazabilidad."
    )

    add_metadata_table(
        doc,
        [
            ("Versión", "1.0"),
            ("Estado", "Documento base para diseño y desarrollo"),
            ("Plataforma recomendada", "Web app responsive / PWA con Supabase"),
            ("Usuarios", "Vecinos, tendero y domiciliario"),
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("1. Contexto y problema a resolver", level=1)
    add_bullets(
        doc,
        [
            "La tienda del conjunto recibe pedidos por llamada y por medios informales, lo que genera confusión sobre qué se pidió, quién lo pidió y cómo se pagó.",
            "Los comprobantes de transferencia se han vuelto una prueba manual, frágil y fácil de perder.",
            "El tendero necesita registrar ventas, controlar inventario y manejar promociones sin depender de cuadernos o memorias dispersas.",
            "El domiciliario necesita una cola clara de pedidos aprobados, con dirección precisa y estado confiable.",
            "La app debe servir tanto para el conjunto principal como para conjuntos cercanos, sin perder control operativo.",
        ],
    )

    doc.add_heading("2. Objetivo del producto", level=1)
    add_callout(
        doc,
        "Objetivo principal:",
        "Permitir que el cliente elija productos desde un catálogo, pague de forma controlada, y que el tendero apruebe, cobre, despache y registre todo con trazabilidad y mínima fricción."
    )
    add_bullets(
        doc,
        [
            "Prioridad 1: que el cliente pueda seleccionar productos con facilidad y completar pedidos sin confusión.",
            "Prioridad 2: que el tendero tenga seguimiento claro de pagos e inventario.",
            "Prioridad 3: que el domiciliario reciba solo pedidos aprobados y con destino claro.",
            "Prioridad 4: que la tienda pueda crecer con promociones, anuncios, categorías dinámicas y futuras mejoras.",
        ],
    )

    doc.add_heading("3. Visión de producto", level=1)
    add_bullets(
        doc,
        [
            "Sensación visual: premium, moderna, limpia y cercana.",
            "Tono de marca: confiable, práctico y comunitario.",
            "Experiencia: pocos pasos, lectura rápida, botones claros y estructura fácil de entender en móvil.",
            "Modelo operativo: una sola app con roles separados y permisos distintos.",
        ],
    )

    doc.add_heading("4. Usuarios, roles y permisos", level=1)
    add_table(
        doc,
        ["Rol", "Responsabilidades", "Permisos principales"],
        [
            ("Vecino / cliente", "Explora catálogo, arma carrito, pregunta por chat, paga, consulta estado y recibe historial.", "Ver catálogo, crear pedidos, subir comprobante, ver estado, ver historial propio."),
            ("Tendero", "Administra productos, aprueba pedidos, registra pagos físicos, controla inventario, promociones y anuncios.", "CRUD de productos/categorías, aprobar/rechazar pedidos, registrar pagos, ver reportes y anuncios."),
            ("Domiciliario", "Consulta pedidos aprobados, entrega en apartamentos o conjuntos cercanos y actualiza estados.", "Ver pedidos listos, ver direcciones autorizadas, cambiar estados de entrega, añadir observaciones."),
        ],
        [Inches(1.15), Inches(2.8), Inches(2.55)],
        header_fill="E8EEF5",
    )

    doc.add_heading("5. Alcance funcional", level=1)
    add_bullets(
        doc,
        [
            "Catálogo de productos con categorías editables por el tendero.",
            "Carrito de compras y flujo de checkout simple.",
            "Chat o canal de consulta para preguntas puntuales sobre productos.",
            "Pago por transferencia con subida de comprobante.",
            "Pago en efectivo con cálculo de cambio.",
            "Aprobación manual del tendero antes de que el domiciliario reciba el pedido.",
            "Seguimiento de inventario, ventas físicas y ventas digitales.",
            "Promociones, descuentos y ventanas emergentes configurables.",
            "Registro opcional de usuarios para pedidos rápidos y historial.",
            "Entrega dentro del conjunto y a conjuntos cercanos.",
        ],
    )

    doc.add_heading("6. Flujo principal de compra", level=1)
    add_numbers(
        doc,
        [
            "El cliente entra a la app y ve un anuncio opcional, si el tendero lo activó.",
            "Explora el catálogo por categorías o busca un producto específico.",
            "Agrega productos al carrito y revisa cantidades, precios y subtotal.",
            "Elige si desea comprar como invitado o iniciar sesión para guardar datos.",
            "Confirma dirección, apartamento, torre o conjunto cercano, según aplique.",
            "Selecciona método de pago: transferencia o efectivo.",
            "Si paga por transferencia, sube una foto del comprobante.",
            "Si paga en efectivo, indica con cuánto paga para calcular el cambio.",
            "El tendero revisa el pedido y lo aprueba o rechaza con motivo.",
            "Si se aprueba, el domiciliario lo recibe en su cola de entregas.",
            "El domiciliario entrega el pedido y marca el estado final.",
        ],
    )

    doc.add_heading("7. Reglas de negocio de pagos y aprobación", level=1)
    add_bullets(
        doc,
        [
            "Un pedido con transferencia no pasa a despacho hasta que el tendero lo apruebe.",
            "La foto de transferencia se guarda como evidencia asociada al pedido.",
            "El tendero puede registrar pagos físicos manualmente para mantener una contabilidad limpia.",
            "Si el cliente paga en efectivo, el sistema debe mostrar el cambio estimado y dejarlo visible en el pedido.",
            "El domiciliario nunca debe ver pedidos pendientes de revisión ni pedidos rechazados.",
            "Toda aprobación, rechazo o modificación importante debe quedar en un historial de auditoría.",
        ],
    )

    doc.add_heading("8. Personalización para el tendero", level=1)
    add_table(
        doc,
        ["Área personalizable", "Qué debe poder configurar", "Valor para el negocio"],
        [
            ("Marca", "Nombre visible, logo, colores, imagen principal y tono de mensajes.", "Hace que la app se sienta propia y coherente con el negocio."),
            ("Catálogo", "Crear, editar, ocultar y reordenar categorías y productos.", "Permite adaptarse a cambios de inventario y temporadas."),
            ("Promociones", "Descuentos, combos, productos destacados y reglas de visibilidad.", "Aumenta ventas sin depender de administración técnica."),
            ("Anuncios", "Modal al abrir la app con texto, imagen o ambos; activable y desactivable.", "Sirve para comunicar novedades, horarios o promociones."),
            ("Operación", "Horarios, zonas de entrega, mínimo de pedido y métodos de pago habilitados.", "Evita pedidos fuera de horario o fuera de cobertura."),
            ("Inventario", "Stock inicial, alertas de bajo inventario y ajuste manual de existencias.", "Reduce errores y rupturas de stock."),
        ],
        [Inches(1.55), Inches(2.65), Inches(2.3)],
        header_fill="E8EEF5",
    )

    doc.add_heading("9. Módulos funcionales", level=1)
    add_table(
        doc,
        ["Módulo", "Funcionalidad mínima", "Observaciones"],
        [
            ("Inicio", "Acceso rápido a catálogo, buscador, promociones y estado del pedido.", "Debe cargar rápido en móvil y mostrar la acción principal."),
            ("Catálogo", "Listado por categorías, filtros, imágenes, precio y disponibilidad.", "Las categorías deben ser editables por el tendero."),
            ("Carrito y checkout", "Modificar cantidades, aplicar notas y elegir medio de pago.", "Debe ser lo más corto posible."),
            ("Pedidos", "Estados, historial, detalle, comprobantes y trazabilidad.", "Es el núcleo de confianza de la app."),
            ("Chat", "Consulta puntual sobre productos o disponibilidad.", "Puede ser chat interno o bandeja de preguntas al tendero."),
            ("Inventario", "Stock, ajustes, entradas, salidas y alertas.", "Debe servir para ventas digitales y físicas."),
            ("Panel tendero", "Aprobación, reportes, pagos, categorías, promociones y anuncios.", "Debe ser simple y muy visible en estados."),
            ("Panel domiciliario", "Pedidos aprobados, ruta/ubicación, notas de entrega y confirmación final.", "Debe evitar confusión entre apartamentos y conjuntos."),
        ],
        [Inches(1.3), Inches(3.25), Inches(1.95)],
        header_fill="E8EEF5",
    )

    doc.add_heading("10. Casos de error y manejo esperado", level=1)
    add_table(
        doc,
        ["Escenario", "Causa probable", "Respuesta de la app"],
        [
            ("Comprobante ilegible", "Foto borrosa, oscura o incompleta.", "Solicitar nueva imagen y dejar el pedido en revisión."),
            ("Pago duplicado", "El cliente reenvía evidencia o crea dos pedidos iguales.", "Mostrar alerta al tendero y marcar posible duplicidad."),
            ("Sin stock", "El artículo se agotó después de agregarse al carrito.", "Bloquear confirmación y ofrecer ajuste del carrito."),
            ("Cambio insuficiente", "El cliente indicó una cantidad menor a la necesaria.", "Avisar antes de enviar el pedido y permitir corrección."),
            ("Dirección ambigua", "Apartamento, torre o conjunto mal escritos.", "Forzar validación y pedir confirmación antes de aprobar."),
            ("Pedido no aprobado", "El tendero aún no revisa o rechaza el pedido.", "Mantener estado pendiente o rechazado, sin despacho."),
            ("Imagen no sube", "Fallo de red, tamaño excesivo o formato inválido.", "Permitir reintento y guardar borrador local si es posible."),
            ("Usuario sin sesión", "El cliente intenta guardar historial sin registrarse.", "Permitir modo invitado pero sugerir registro opcional."),
            ("Conexión lenta", "Red móvil inestable en el conjunto.", "Usar carga progresiva, skeletons y reintento automático."),
            ("Conflicto de inventario", "Dos pedidos reducen el mismo stock al mismo tiempo.", "Aplicar control transaccional y recalcular disponibilidad."),
        ],
        [Inches(1.55), Inches(2.15), Inches(2.8)],
        header_fill="F4F6F9",
    )

    doc.add_heading("11. Modelo de datos recomendado en Supabase", level=1)
    add_callout(
        doc,
        "Arquitectura recomendada:",
        "Usar Supabase como backend principal con Auth, Postgres, Storage, Realtime y Edge Functions. El cliente nunca debe recibir credenciales privilegiadas; todo acceso sensible debe protegerse con Row Level Security y funciones del servidor."
    )
    add_table(
        doc,
        ["Entidad", "Propósito", "Campos clave sugeridos"],
        [
            ("profiles", "Perfil y rol de cada usuario.", "id, auth_user_id, role, nombre, apartamento, torre, telefono, activo"),
            ("stores", "Configuración del negocio.", "id, nombre, logo_url, colores, horario, zonas_entrega, anuncio_activo"),
            ("categories", "Categorías editables.", "id, store_id, nombre, orden, visible"),
            ("products", "Catálogo de productos.", "id, store_id, category_id, nombre, precio, stock, imagen_url, activo"),
            ("orders", "Pedido principal.", "id, customer_id, store_id, estado, metodo_pago, total, direccion, notas"),
            ("order_items", "Detalle de productos por pedido.", "id, order_id, product_id, cantidad, precio_unitario"),
            ("payments", "Registro de pagos y evidencias.", "id, order_id, metodo, estado, comprobante_url, cash_received, cash_change"),
            ("deliveries", "Estado logístico del domiciliario.", "id, order_id, rider_id, estado, confirmacion_entrega, observaciones"),
            ("announcements", "Banner o modal promocional.", "id, store_id, titulo, texto, imagen_url, activo, tipo"),
            ("audit_logs", "Registro de acciones críticas.", "id, actor_id, accion, entidad, entidad_id, created_at"),
        ],
        [Inches(1.35), Inches(1.95), Inches(3.2)],
        header_fill="E8EEF5",
    )

    doc.add_heading("12. Requisitos de seguridad", level=1)
    add_table(
        doc,
        ["Riesgo", "Control de seguridad", "Implementación recomendada"],
        [
            ("Acceso indebido a pedidos", "Autorización por rol y RLS.", "Cada tabla sensible debe tener políticas que filtren por usuario, tienda y rol."),
            ("Manipulación de comprobantes", "Storage privado y URLs firmadas.", "Las evidencias no deben quedar públicas."),
            ("Uso del service role en cliente", "Secreto solo del servidor.", "Nunca exponer service role en frontend."),
            ("Aprobaciones falsas", "Acciones críticas server-side.", "Aprobar/rechazar pedidos solo desde funciones seguras."),
            ("Spam o abuso en chat", "Rate limiting y moderación.", "Limitar frecuencia y registrar incidentes."),
            ("Datos sensibles expuestos", "Principio de mínimo privilegio.", "Mostrar solo lo necesario por rol."),
            ("Sesiones robadas", "MFA opcional para tendero, expiración y revocación.", "Activar autenticación robusta en cuentas administrativas."),
            ("Subidas maliciosas", "Validación de tamaño y tipo de archivo.", "Permitir solo formatos esperados y límites claros."),
        ],
        [Inches(1.6), Inches(2.15), Inches(2.75)],
        header_fill="F4F6F9",
    )
    add_bullets(
        doc,
        [
            "Usar HTTPS en todo momento.",
            "Registrar logs de auditoría para cambios de precio, inventario, aprobación de pedidos y cancelaciones.",
            "Aplicar validación tanto en frontend como en backend.",
            "Separar claramente permisos de cliente, tendero y domiciliario.",
            "Habilitar políticas de backup y restauración de la base de datos.",
        ],
    )

    doc.add_heading("13. Responsive, accesibilidad y experiencia móvil", level=1)
    add_bullets(
        doc,
        [
            "Diseño mobile-first con adaptación a tablets y escritorio.",
            "Botones táctiles con altura mínima de 44 px.",
            "Contraste suficiente entre texto, fondo y estados de acción.",
            "Estados visibles de foco, error, éxito, deshabilitado y carga.",
            "Evitar tablas o bloques que obliguen a hacer zoom horizontal en móvil.",
            "Priorizar componentes apilados en pantallas pequeñas y rejillas en pantallas grandes.",
            "Usar imágenes optimizadas y carga diferida para no castigar datos móviles.",
        ],
    )

    doc.add_heading("14. Pantallas recomendadas", level=1)
    add_numbers(
        doc,
        [
            "Inicio / portada.",
            "Catálogo y búsqueda.",
            "Detalle de producto.",
            "Carrito.",
            "Checkout y pago.",
            "Historial y estado del pedido.",
            "Chat o consultas.",
            "Panel del tendero.",
            "Panel del domiciliario.",
            "Ajustes y personalización de la tienda.",
        ],
    )

    doc.add_heading("15. Criterios de aceptación del MVP", level=1)
    add_bullets(
        doc,
        [
            "Un vecino puede crear un pedido completo sin apoyo manual.",
            "El tendero puede aprobar o rechazar un pedido con evidencia visible.",
            "El domiciliario solo ve pedidos aprobados.",
            "El tendero puede editar productos, categorías, inventario y promociones.",
            "La app guarda historial de pedidos y pagos.",
            "La app registra pagos físicos y transferencias con trazabilidad.",
            "La experiencia funciona correctamente en móvil y escritorio.",
            "Las reglas de seguridad impiden ver datos ajenos a cada rol.",
        ],
    )

    doc.add_heading("16. Recomendación técnica final", level=1)
    add_callout(
        doc,
        "Recomendación de ingeniería:",
        "Construir la solución como una PWA responsive con una sola base de código, backend en Supabase y paneles por rol. Esa combinación da velocidad de desarrollo, buena experiencia móvil, control de acceso fino y una evolución ordenada hacia funciones más avanzadas."
    )
    add_bullets(
        doc,
        [
            "Frontend sugerido: React / Next.js con TypeScript y UI mobile-first.",
            "Backend sugerido: Supabase Auth, Postgres, Storage, Realtime y Edge Functions.",
            "Despliegue sugerido: hosting con HTTPS, variables seguras y monitoreo básico.",
            "Evolución futura: notificaciones push, reportes avanzados, fidelización y soporte para varias tiendas.",
        ],
    )

    doc.add_paragraph("")
    p_end = doc.add_paragraph()
    r_end = p_end.add_run(
        "Este documento está preparado como base de construcción del producto. El siguiente paso ideal es convertirlo en backlog, arquitectura de datos y plan de implementación por sprints."
    )
    fmt_run(r_end, size=10.5, color=MUTED)

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUT)
    print(OUT)
