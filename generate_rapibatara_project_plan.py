from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_rapibatara_spec import (
    DARK,
    MUTED,
    add_bullets,
    add_callout,
    add_metadata_table,
    add_numbers,
    add_table,
    fmt_run,
    format_doc,
)


OUT = Path(r"C:\Users\pipev\OneDrive\Documentos\RapiBatará\rapibatara_plan_proyecto_v1.docx")


def add_section_intro(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    fmt_run(r, size=10.5, color=DARK)


def build_document():
    doc = Document()
    format_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("RapiBatará")
    fmt_run(r, size=26, color="000000", bold=False)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("Plan de proyecto para arrancar desarrollo")
    fmt_run(r2, size=11, color=MUTED)

    add_callout(
        doc,
        "Propósito:",
        "Definir el plan práctico para iniciar la construcción de RapiBatará con una base sólida: stack, arquitectura, organización del trabajo, seguridad, despliegue y secuencia de sprints."
    )

    add_metadata_table(
        doc,
        [
            ("Versión", "1.0"),
            ("Objetivo", "Arrancar desarrollo con orden y bajo riesgo"),
            ("Modelo", "PWA responsive + Supabase"),
            ("Enfoque", "MVP operativo primero, mejoras después"),
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("1. Resultado esperado del proyecto", level=1)
    add_bullets(
        doc,
        [
            "Un vecino puede registrarse o entrar como invitado, hacer un pedido y pagar sin confusión.",
            "El tendero puede administrar catálogo, pedidos, inventario, cobros y anuncios desde un panel único.",
            "El domiciliario puede ver solo pedidos aprobados y entregar con claridad.",
            "Toda acción importante queda trazada para evitar discusiones por pagos, entregas o stock.",
        ],
    )

    doc.add_heading("2. Alcance del MVP", level=1)
    add_table(
        doc,
        ["Incluye", "No incluye en el primer release"],
        [
            ("Catálogo, carrito, checkout, comprobante de transferencia, pago en efectivo, aprobación, inventario, panel tendero, panel domiciliario y auditoría.", "App nativa iOS/Android separada, multi-tienda compleja, programas de fidelización avanzados, analítica pesada, automatizaciones complejas."),
        ],
        [Inches(3.2), Inches(3.3)],
        header_fill="E8EEF5",
    )

    doc.add_heading("3. Stack recomendado", level=1)
    add_table(
        doc,
        ["Capa", "Recomendación", "Motivo"],
        [
            ("Frontend", "Next.js + React + TypeScript", "Buen equilibrio entre velocidad, mantenibilidad y experiencia responsive."),
            ("UI", "Componentes propios o shadcn/ui", "Permite una interfaz limpia, rápida de construir y fácil de personalizar."),
            ("Backend", "Supabase", "Auth, Postgres, Storage, Realtime y Edge Functions en una sola plataforma."),
            ("Estado", "Server components + estado local ligero", "Reduce complejidad y mantiene la app fluida."),
            ("Validación", "Zod o esquema equivalente", "Evita entradas inválidas desde cliente y servidor."),
            ("Deploy", "Vercel o plataforma similar", "Entrega rápida, HTTPS y despliegue continuo."),
        ],
        [Inches(1.2), Inches(2.5), Inches(2.8)],
        header_fill="F4F6F9",
    )

    doc.add_heading("4. Arquitectura de alto nivel", level=1)
    add_bullets(
        doc,
        [
            "Aplicación web responsive pensada primero para móvil.",
            "Supabase como backend único para autenticación, datos y archivos.",
            "Acciones sensibles ejecutadas en Edge Functions o rutas del servidor.",
            "RLS obligatoria en todas las tablas de datos sensibles.",
            "Storage privado para comprobantes, logos e imágenes administradas.",
            "Auditoría de eventos para cambios de precios, inventario, pagos y estados.",
        ],
    )

    doc.add_heading("5. Modelo de módulos", level=1)
    add_table(
        doc,
        ["Módulo", "Responsable funcional", "Entrega mínima"],
        [
            ("Acceso y perfiles", "Auth y roles", "Login, registro opcional, invitado, perfiles por rol."),
            ("Catálogo", "Producto/tienda", "Listado, categorías, búsqueda, detalle."),
            ("Checkout", "Compra", "Carrito, pago, comprobante, efectivo."),
            ("Pedidos", "Operación", "Estados, historial, cancelación controlada."),
            ("Inventario", "Tendero", "Stock, alertas, ajustes, entradas y salidas."),
            ("Despacho", "Domiciliario", "Pedidos aprobados, direcciones, confirmación."),
            ("Promociones", "Tendero", "Anuncios, descuentos, modal de entrada."),
            ("Auditoría", "Sistema", "Registro de acciones críticas."),
        ],
        [Inches(1.35), Inches(1.9), Inches(3.25)],
        header_fill="E8EEF5",
    )

    doc.add_heading("6. Estructura de trabajo recomendada", level=1)
    add_table(
        doc,
        ["Área", "Responsabilidad", "Salida"],
        [
            ("Producto", "Definir prioridades, validar reglas, aprobar alcance.", "Backlog vivo y decisiones cerradas."),
            ("Diseño", "Sistema visual, componentes, estados y responsive.", "UI kit y pantallas clave."),
            ("Frontend", "Implementar interfaz y flujos.", "App funcional conectada al backend."),
            ("Backend", "Tablas, políticas, funciones, storage y auth.", "Base segura y consistente."),
            ("QA", "Pruebas funcionales, seguridad y regresión.", "Lista de bugs y validación de release."),
        ],
        [Inches(1.2), Inches(2.8), Inches(2.5)],
        header_fill="F4F6F9",
    )

    doc.add_heading("7. Diseño del repositorio", level=1)
    add_table(
        doc,
        ["Carpeta", "Contenido esperado"],
        [
            ("/app", "Rutas y pantallas del producto."),
            ("/components", "Componentes reutilizables de UI."),
            ("/features", "Lógica por dominio: catálogo, pedidos, pagos, inventario."),
            ("/lib", "Clientes, helpers, validación y utilidades."),
            ("/supabase", "Migraciones, seeds y políticas."),
            ("/public", "Assets públicos como logo e imágenes base."),
            ("/docs", "Decisiones, backlog, criterios y notas técnicas."),
        ],
        [Inches(1.8), Inches(4.7)],
        header_fill="F4F6F9",
    )

    doc.add_heading("8. Plan de sprints", level=1)
    add_table(
        doc,
        ["Sprint", "Objetivo", "Entregable"],
        [
            ("Sprint 0", "Preparar base técnica y decisiones cerradas.", "Repo, Supabase, estructura y diseño inicial."),
            ("Sprint 1", "Autenticación, roles y perfiles.", "Acceso funcional para vecino, tendero y domiciliario."),
            ("Sprint 2", "Catálogo, categorías y carrito.", "Exploración y armado de pedido."),
            ("Sprint 3", "Checkout y pagos.", "Transferencia con comprobante y efectivo con cambio."),
            ("Sprint 4", "Aprobación y despacho.", "Tendero aprueba, domiciliario entrega."),
            ("Sprint 5", "Inventario, cobros físicos y auditoría.", "Control operativo real de tienda."),
            ("Sprint 6", "Promociones, anuncios y pulido UX.", "Panel administrativo más completo."),
            ("Sprint 7", "QA, seguridad, performance y salida.", "Versión lista para uso real."),
        ],
        [Inches(1.1), Inches(3.2), Inches(2.2)],
        header_fill="E8EEF5",
    )

    doc.add_heading("9. Secuencia técnica de implementación", level=1)
    add_numbers(
        doc,
        [
            "Crear el proyecto web y configurar el entorno de desarrollo.",
            "Conectar Supabase y definir autenticación, roles y políticas iniciales.",
            "Modelar las tablas y relaciones principales.",
            "Construir la base visual y los componentes reutilizables.",
            "Implementar catálogo, carrito y checkout.",
            "Agregar comprobantes, pagos y aprobación.",
            "Conectar inventario, pedidos y despacho.",
            "Sumar promociones, anuncios y personalización.",
            "Ejecutar QA, auditoría y hardening de seguridad.",
            "Preparar despliegue, monitoreo y documentación de uso.",
        ],
    )

    doc.add_heading("10. Reglas de seguridad desde el día 1", level=1)
    add_bullets(
        doc,
        [
            "Nunca exponer service role key en cliente.",
            "Usar RLS en todas las tablas sensibles.",
            "Guardar comprobantes en Storage privado.",
            "Limitar acceso por rol y por tienda.",
            "Registrar auditoría en todas las acciones críticas.",
            "Validar archivos por tamaño, tipo y origen.",
            "Proteger acciones de aprobación y ajustes manuales en servidor.",
            "Aplicar sesiones seguras y expiración razonable.",
        ],
    )

    doc.add_heading("11. Estrategia de pruebas", level=1)
    add_table(
        doc,
        ["Tipo de prueba", "Qué valida", "Ejemplos"],
        [
            ("Funcional", "Que cada flujo haga lo que promete.", "Crear pedido, aprobar, entregar, registrar pago."),
            ("Permisos", "Que cada rol vea lo suyo.", "Cliente no ve panel del tendero."),
            ("Datos", "Que no haya inconsistencias.", "Stock no negativo, totales correctos."),
            ("Responsive", "Que funcione bien en móvil y escritorio.", "Checkout usable en pantallas pequeñas."),
            ("Seguridad", "Que no haya fugas ni rutas abiertas.", "Comprobantes privados y RLS activada."),
            ("Regresión", "Que cambios no dañen funciones previas.", "Editar inventario sin romper pedidos."),
        ],
        [Inches(1.4), Inches(2.4), Inches(2.7)],
        header_fill="F4F6F9",
    )

    doc.add_heading("12. Definición de listo para arrancar", level=1)
    add_bullets(
        doc,
        [
            "La especificación técnica y el backlog están aprobados.",
            "El stack está decidido.",
            "Las tablas y roles principales están definidos.",
            "El MVP está delimitado.",
            "Existe una secuencia de sprints.",
            "Se conocen riesgos, seguridad y pruebas mínimas.",
        ],
    )

    doc.add_heading("13. Riesgos y mitigaciones", level=1)
    add_table(
        doc,
        ["Riesgo", "Mitigación"],
        [
            ("Cambiar alcance a mitad del desarrollo", "Congelar MVP y usar backlog de mejoras posteriores."),
            ("Sobrecargar la primera versión", "Construir solo lo necesario para vender, cobrar y despachar."),
            ("Problemas de permisos o datos", "RLS, pruebas de roles y acciones críticas en servidor."),
            ("Fallo de adopción por parte del tendero", "Pantallas simples, flujos cortos y validación temprana con el negocio."),
            ("Lentitud en móvil", "Optimización de assets, carga diferida y UI liviana."),
        ],
        [Inches(2.5), Inches(4.0)],
        header_fill="F4F6F9",
    )

    doc.add_heading("14. Cierre ejecutivo", level=1)
    add_callout(
        doc,
        "Decisión recomendada:",
        "Empezar con Sprint 0 y Sprint 1 de inmediato: base técnica, Supabase, roles, autenticación y estructura. En paralelo, dejar congeladas las reglas de negocio de pago, aprobación y despacho para que el desarrollo avance sin improvisación."
    )

    doc.add_paragraph("")
    p_end = doc.add_paragraph()
    r_end = p_end.add_run(
        "Este plan ya está suficientemente detallado para pasar a ejecución. El siguiente documento natural es el setup técnico del proyecto: estructura del repo, esquema de Supabase, políticas RLS y primeras pantallas."
    )
    fmt_run(r_end, size=10.5, color=MUTED)

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUT)
    print(OUT)
