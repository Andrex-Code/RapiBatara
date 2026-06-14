from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_rapibatara_spec import (
    ACCENT,
    DARK,
    DARK_BLUE,
    LIGHT,
    MUTED,
    add_bullets,
    add_callout,
    add_metadata_table,
    add_numbers,
    add_para,
    add_table,
    fmt_run,
    format_doc,
)


OUT = Path(r"C:\Users\pipev\OneDrive\Documentos\RapiBatará\rapibatara_backlog_v1.docx")


def add_section_intro(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    fmt_run(r, size=10.5, color=DARK)


def build_document():
    doc = Document()
    format_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("RapiBatará")
    fmt_run(r, size=26, color="000000", bold=False)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("Backlog funcional, roadmap y plan de implementación")
    fmt_run(r2, size=11, color=MUTED)

    add_callout(
        doc,
        "Propósito del documento:",
        "Convertir la especificación técnica en trabajo ejecutable para diseño, desarrollo, pruebas y salida a producción. La prioridad es construir un MVP confiable y luego extenderlo sin romper la operación diaria del tendero."
    )

    add_metadata_table(
        doc,
        [
            ("Versión", "1.0"),
            ("Enfoque", "MVP + base para evolución"),
            ("Prioridad", "Cobrar, vender, aprobar, despachar y auditar"),
            ("Supabase", "Auth + Postgres + Storage + Realtime + Edge Functions"),
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("1. Principios de priorización", level=1)
    add_bullets(
        doc,
        [
            "Primero resolver el problema real de pedidos y pagos, no la estética secundaria.",
            "Cada historia debe reducir fricción para el cliente o carga operativa para el tendero.",
            "Las funciones de administración deben ser más importantes que los adornos de marketing.",
            "La seguridad y trazabilidad no son extras: son parte del producto base.",
            "Todo lo que afecte dinero, inventario o despacho requiere control de permisos y auditoría.",
        ],
    )

    doc.add_heading("2. Epics principales", level=1)
    add_table(
        doc,
        ["Epic", "Objetivo", "Prioridad"],
        [
            ("E1. Acceso y perfiles", "Autenticación, roles y perfiles por usuario.", "Alta"),
            ("E2. Catálogo y carrito", "Explorar, buscar, agregar y revisar productos.", "Alta"),
            ("E3. Checkout y pagos", "Transferencia, efectivo, comprobantes y estados.", "Alta"),
            ("E4. Aprobación y despacho", "Validación del tendero y entrega del domiciliario.", "Alta"),
            ("E5. Inventario y ventas", "Control operativo del negocio.", "Alta"),
            ("E6. Promociones y anuncios", "Descuentos, banners y mensajes al abrir.", "Media"),
            ("E7. Chat y soporte", "Consultas puntuales del cliente.", "Media"),
            ("E8. Reportes y auditoría", "Historial, trazabilidad y control.", "Alta"),
            ("E9. Personalización de tienda", "Marca, horarios, zonas y preferencias.", "Media"),
            ("E10. Notificaciones y mejoras futuras", "Push, automatizaciones y expansión.", "Baja para MVP"),
        ],
        [Inches(1.6), Inches(3.6), Inches(1.3)],
        header_fill="E8EEF5",
    )

    doc.add_heading("3. Roadmap sugerido por fases", level=1)
    add_table(
        doc,
        ["Fase", "Alcance", "Resultado esperado"],
        [
            ("Fase 0. Descubrimiento", "Validar flujos, entidades, reglas de negocio y alcance real.", "Documento final + decisiones cerradas."),
            ("Fase 1. Base técnica", "Supabase, autenticación, estructura de datos y roles.", "Usuarios entrando con permisos correctos."),
            ("Fase 2. Compra", "Catálogo, carrito, checkout y medios de pago.", "Un cliente puede hacer un pedido completo."),
            ("Fase 3. Operación", "Aprobación, inventario, pagos físicos y despacho.", "El tendero y domiciliario operan sin cuaderno."),
            ("Fase 4. Administración", "Promos, anuncios, configuración y auditoría.", "La tienda puede administrarse sola en el día a día."),
            ("Fase 5. Optimización", "Analytics, notificaciones, mejoras de UX y performance.", "Sistema preparado para escalar y crecer."),
        ],
        [Inches(1.35), Inches(3.45), Inches(1.85)],
        header_fill="E8EEF5",
    )

    doc.add_heading("4. Backlog priorizado del MVP", level=1)
    add_table(
        doc,
        ["ID", "Historia de usuario", "Criterio de aceptación", "Prioridad"],
        [
            ("US-01", "Como vecino, quiero entrar a la app y ver el catálogo sin registrarme.", "Puedo navegar productos en modo invitado.", "Alta"),
            ("US-02", "Como vecino, quiero crear una cuenta opcional para guardar pedidos y datos.", "Puedo registrarme y volver a entrar con mis datos.", "Alta"),
            ("US-03", "Como vecino, quiero buscar productos y filtrarlos por categoría.", "Encuentro productos rápido en menos toques.", "Alta"),
            ("US-04", "Como vecino, quiero agregar productos al carrito y editar cantidades.", "El carrito recalcula totales correctamente.", "Alta"),
            ("US-05", "Como vecino, quiero dejar notas para mi pedido.", "La nota se guarda en el pedido.", "Media"),
            ("US-06", "Como vecino, quiero pagar por transferencia y subir comprobante.", "La imagen se adjunta y el pedido queda pendiente.", "Alta"),
            ("US-07", "Como vecino, quiero pagar en efectivo e indicar con cuánto pago.", "El sistema calcula el cambio estimado.", "Alta"),
            ("US-08", "Como vecino, quiero ver el estado de mi pedido.", "Veo pendiente, aprobado, en camino y entregado.", "Alta"),
            ("US-09", "Como tendero, quiero aprobar o rechazar pedidos.", "El pedido cambia de estado y queda auditado.", "Alta"),
            ("US-10", "Como tendero, quiero registrar pagos físicos manualmente.", "Puedo sumar ventas fuera de la app con trazabilidad.", "Alta"),
            ("US-11", "Como tendero, quiero crear y editar productos.", "El catálogo refleja cambios sin intervención técnica.", "Alta"),
            ("US-12", "Como tendero, quiero crear y editar categorías.", "Puedo reorganizar el catálogo libremente.", "Alta"),
            ("US-13", "Como tendero, quiero ajustar inventario.", "El stock se actualiza y bloquea ventas sin disponibilidad.", "Alta"),
            ("US-14", "Como domiciliario, quiero ver solo pedidos aprobados.", "No veo pedidos pendientes o rechazados.", "Alta"),
            ("US-15", "Como domiciliario, quiero ver a qué apartamento o conjunto debo ir.", "La dirección está clara y validada.", "Alta"),
            ("US-16", "Como tendero, quiero publicar anuncios promocionales.", "Puedo activar/desactivar un banner o modal.", "Media"),
        ],
        [Inches(0.7), Inches(3.0), Inches(2.1), Inches(0.7)],
        header_fill="F4F6F9",
    )

    doc.add_heading("5. Backlog técnico", level=1)
    add_table(
        doc,
        ["ID", "Tarea técnica", "Motivo", "Prioridad"],
        [
            ("TECH-01", "Configurar Supabase Auth con roles.", "Base de seguridad y experiencia de acceso.", "Alta"),
            ("TECH-02", "Definir esquema de Postgres y relaciones.", "Evita inconsistencias desde el inicio.", "Alta"),
            ("TECH-03", "Implementar RLS en todas las tablas sensibles.", "Protege datos por usuario y rol.", "Alta"),
            ("TECH-04", "Configurar Storage privado para comprobantes.", "La evidencia de pago no debe ser pública.", "Alta"),
            ("TECH-05", "Crear Edge Functions para aprobaciones y acciones críticas.", "Evita lógica sensible en frontend.", "Alta"),
            ("TECH-06", "Definir auditoría de acciones críticas.", "Necesaria para pagos, inventario y cambios.", "Alta"),
            ("TECH-07", "Preparar estructura responsive mobile-first.", "La app se usará principalmente en móvil.", "Alta"),
            ("TECH-08", "Agregar manejo de errores y estados de carga.", "Reduce confusión y mejora confianza.", "Alta"),
            ("TECH-09", "Preparar observabilidad básica.", "Ayuda a diagnosticar fallos reales.", "Media"),
            ("TECH-10", "Preparar notificaciones futuras.", "Deja lista la base para iteraciones posteriores.", "Baja"),
        ],
        [Inches(0.8), Inches(2.7), Inches(2.2), Inches(0.8)],
        header_fill="F4F6F9",
    )

    doc.add_heading("6. Dependencias críticas", level=1)
    add_bullets(
        doc,
        [
            "La autenticación debe existir antes de activar permisos por rol.",
            "El esquema de pedidos debe existir antes del flujo de checkout.",
            "El Storage privado debe estar listo antes de aceptar comprobantes de transferencia.",
            "Las políticas RLS deben probarse antes de abrir acceso real a usuarios.",
            "El panel del tendero depende del modelo de datos de productos, inventario y pedidos.",
            "El panel del domiciliario depende del estado de aprobación del pedido.",
        ],
    )

    doc.add_heading("7. Riesgos de entrega", level=1)
    add_table(
        doc,
        ["Riesgo", "Impacto", "Mitigación"],
        [
            ("Complejidad excesiva del primer release", "Retrasa salida a producción.", "Limitar MVP a compra, aprobación, inventario y cobro."),
            ("Errores de permisos", "Filtración de datos o acciones indebidas.", "RLS estricta, pruebas de rol y funciones del servidor."),
            ("Cambios improvisados de alcance", "Rompen el orden del roadmap.", "Congelar alcance por fase y gestionar cambios explícitamente."),
            ("Falta de validación operativa", "La app no encaja con la tienda real.", "Validar flujos con el tendero antes de construir la interfaz final."),
            ("Poca calidad de red móvil", "Mala experiencia en conjuntos.", "Diseño liviano, carga diferida y reintentos."),
        ],
        [Inches(1.6), Inches(2.1), Inches(2.8)],
        header_fill="F4F6F9",
    )

    doc.add_heading("8. Definición de listo para pasar a proyecto", level=1)
    add_numbers(
        doc,
        [
            "La especificación funcional está cerrada y aprobada.",
            "Los roles y permisos están definidos sin ambigüedades.",
            "El MVP está limitado a funcionalidades de alto valor.",
            "El modelo de datos principal está definido.",
            "Las reglas de seguridad y auditoría están consideradas.",
            "El backlog contiene historias suficientes para iniciar sprint 1.",
        ],
    )

    doc.add_heading("9. Próximo paso recomendado", level=1)
    add_callout(
        doc,
        "Siguiente entrega sugerida:",
        "Crear el plan de proyecto inicial con stack, arquitectura de carpetas, modelo de datos en Supabase, estructura de componentes y lista de tareas del sprint 1. Ese sería el punto ideal para empezar a construir la app."
    )

    doc.add_paragraph("")
    p_end = doc.add_paragraph()
    r_end = p_end.add_run(
        "Este backlog ya sirve como puente entre la idea de negocio y la ejecución técnica. A partir de aquí, el siguiente documento útil es el plan de proyecto y arquitectura de implementación."
    )
    fmt_run(r_end, size=10.5, color=MUTED)

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUT)
    print(OUT)
